"""
生成项目汇报发言稿 DOCX
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# ---- 页面设置 ----
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.6
style.paragraph_format.space_after = Pt(6)
# 中文字体回退
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ============================================================
# 标题
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('基于AI智能的电商商品销售分析与预测系统')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1F, 0x59, 0x67)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('课堂汇报发言稿（5–6分钟）')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('团队：严辰乐（组长）、苏文韬、姚凯曦、薛淞、闫维岳\n中山大学《软件工程》课程实践 · 选题三(1)')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_paragraph('')

# ---- 辅助函数 ----
def add_section(seq, title_text, duration, content):
    """添加一个发言章节"""
    # 章节标题
    h = doc.add_heading(f'{seq}. {title_text}　　⏱ {duration}', level=2)
    for run in h.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1F, 0x59, 0x67)
    # 正文
    for para_text in content:
        p = doc.add_paragraph(para_text)

# ============================================================
# 各节内容
# ============================================================

add_section(1, '首页封面 — 开场', '约30秒', [
    '各位老师、同学，大家好。',
    '我们小组做的是选题三，基于AI智能的电商商品销售分析与预测系统。简单来说，这是一个面向电商企业内部的商业智能分析平台。它帮运营人员解决一个问题：从海量销售数据中快速看出趋势、发现异常、做好预测，而不是靠人工翻报表。',
    '我们项目一共实现了8个前端页面、23条后端API、10张核心数据表，96项集成测试全部通过。下面我按五个方面来介绍。',
])

add_section(2, '项目背景', '约40秒', [
    '电商平台每天会产生大量的交易、商品、客户和库存数据。传统做法是靠数据分析师手工拉报表，效率低，而且很多问题是事后才发现的——比如某品类销售额突然下滑了30%，可能过了一周才有人注意到。',
    '另一方面，选题说明里也明确要求系统要具备销售预测、库存优化、用户画像、异常检测这些能力。',
    '所以我们希望构建一个平台，把"数据查询 → 趋势预测 → 异常预警 → 报表输出"串成一条完整的分析链路，让运营经理打开浏览器就能掌握经营状态。',
])

add_section(3, '技术方案', '约60秒', [
    '系统采用B/S四层架构：前端展示层、后端业务逻辑层、数据访问层和ML服务层。',
    '前端这块，我们用Vue 3加Element Plus组件库，图表用ECharts渲染。这样仪表盘、查询筛选、预测曲线这些都能在一个统一框架下实现。',
    '后端是Python Flask，用PyJWT做登录认证和角色权限控制，用APScheduler做定时预警扫描。数据库是MySQL，一共设计了10张表来覆盖商品、客户、销售记录、预测结果、预警规则等业务实体。',
    'ML算法层用的是scikit-learn，核心是线性回归模型。我们训练了一个7维特征的预测模型来做未来7天的销量预测，然后结合安全库存公式给出补货建议。异常检测目前用的是规则引擎，基于7日移动均线做基线对比。',
    '整个系统前后端分离，前端通过RESTful API调用后端，统一JSON格式通信。技术方案图、用例图、构件图、ER图在演示页上都有展示，大家可以直观看到。',
])

add_section(4, '软件功能', '约90秒', [
    '系统目前实现了8个前端页面，核心功能围绕6个模块展开。',
    '第一个是登录与权限控制。用户通过账号密码登录，系统根据admin、manager、analyst不同角色控制可见菜单和操作权限。同时我们还做了一个演示模式，不需要启动后端和数据库就能浏览全部页面，方便开发阶段的验证。',
    '第二个是仪表盘首页。用户登录后首先看到总销售额、订单量、毛利率这几张指标卡片，下面是销售趋势图和品类分布图，还有待处理预警的概览。',
    '第三个是多维销售数据查询，这也是P0用例里面最核心的功能。用户可以按时间范围、地区、品类、渠道自由组合筛选，结果以表格和图表并排展示。',
    '第四个是销量预测与库存建议。系统用训练好的线性回归模型预测未来7天每个商品的销量，再结合当前库存库存量计算安全库存和补货建议。',
    '第五个是异常预警。我们预设了两类规则：销售额突降告警和库存警戒线告警。系统每小时自动扫描一次，检测到异常就在预警中心高亮显示，红色是严重、橙色是警告、黄色是提示。运营人员可以查看详情，标记已处理或忽略。',
    '第六个是报表输出。支持生成并下载Excel报表。另外管理后台还有用户管理和数据源配置功能。',
])

add_section(5, '任务分工', '约40秒', [
    '我们团队五个人，按模块化分工来推进。',
    '我作为组长，负责系统架构设计、后端核心开发——包括JWT认证、预测和预警API、定时调度，还有全部文档的主笔和最终审核。',
    '苏文韬负责数据库设计和后端数据层，10张表的DDL、数据查询API、报表导出API，以及1200条随机种子数据的生成脚本。',
    '姚凯曦负责前端管理类页面——登录、预警、用户管理、数据源配置，还有前端整体的API封装层、演示模式的实现，以及96项集成测试和测试文档。',
    '薛淞负责全部机器学习模块，包括销售预测、库存补货建议、异常检测、用户画像和营销效果评估，一共8步ML管道。',
    '闫维岳负责演示材料——这个HTML交互演示页就是他做的，还有7张系统界面截图、UML图的PNG导出，以及前后端联调测试。',
])

add_section(6, '项目管理方式', '约50秒', [
    '项目周期大约两周，我们采用分阶段推进加里程碑控制的方式来管理。',
    '整体分四个阶段：前三天是需求收尾和代码骨架搭建，中间五天集中攻克5个P0核心功能，接下来四天补P1功能加测试修复，最后两天统一整理提交。',
    '我们设了6个里程碑，从M1需求定稿到M6最终提交，每个里程碑都有明确的完成标准。最关键的是第三天的P0铁律——如果第8天P0还没跑通，全体停下手头工作集中攻坚，P1和P2直接砍掉。最终我们按时完成了全部P0和P1功能。',
    '协作方面，我们用GitHub做版本管理，第二天就锁死了API接口文档让前后端并行开发。风险管理上，对数据库延迟、ML精度不达标、前端工作量超预期等7个风险都提前准备了缓解措施。',
    '甘特图直观展示了22个任务在14天内的分布情况，演示页上可以直接看到。',
])

add_section(7, '总结', '约30秒', [
    '总结一下。我们围绕电商运营场景，完成了一个集销售分析、销量预测、库存建议、异常预警和报表导出于一体的BI系统原型。',
    '虽然跟成熟的商业产品还有距离，但这个项目完整走过了需求分析、系统设计、编码实现、集成测试和项目管理的全流程，体现了软件工程不是只写代码，还包括文档、分工、版本协作和风险管控。',
    '后续如果继续迭代，可以接入真实的电商平台数据源，把线性回归模型替换为更先进的时序预测算法，并且完善自动化部署方案。',
    '以上就是我们小组的汇报，谢谢老师和同学们。',
])

# ---- 末尾附注 ----
doc.add_paragraph('')
note = doc.add_paragraph()
run = note.add_run('说明：本发言稿对应 demo/项目演示-交互版.html 的8个section。讲解时可打开HTML演示页，按左侧导航栏顺序逐节配合屏幕切换。总时长约5分20秒，可根据实际情况微调各节节奏。')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
run.italic = True

# ---- 保存 ----
OUTPUT = 'C:/Users/MSI-NB/Desktop/Learning/软工大作业/课堂汇报发言稿.docx'
doc.save(OUTPUT)
print(f'DONE: {OUTPUT}')
